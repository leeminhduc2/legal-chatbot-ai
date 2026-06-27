import json
import sys
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from llm_splitter import LLMVietnameseLegalSplitter


def get_docx_paragraphs(file_path: str | Path) -> List[str]:
    """Đọc file .docx và trả về danh sách các đoạn văn bản (paragraph).
    
    Chỉ lấy các paragraph KHÔNG nằm trong bảng (table).
    Bỏ qua các dòng trống.
    
    Args:
        file_path: Đường dẫn tới file .docx
        
    Returns:
        Danh sách các chuỗi text của từng paragraph.
    """
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def get_docx_tables(file_path: str | Path) -> List[Dict[str, Any]]:
    """Đọc file .docx và trả về danh sách các bảng biểu.
    
    Mỗi bảng được biểu diễn dưới dạng dict gồm:
    - table_index: thứ tự bảng (bắt đầu từ 1)
    - rows: danh sách các hàng, mỗi hàng là list các cell text (đã lọc bỏ ô trống)
    
    Args:
        file_path: Đường dẫn tới file .docx
        
    Returns:
        Danh sách các dict mô tả bảng.
    """
    doc = Document(file_path)
    tables = []
    for index, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            # Lấy text từng ô, lọc bỏ ô trống, loại trùng lặp do merged cells
            seen = set()
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in seen:
                    seen.add(cell_text)
                    row_cells.append(cell_text)
            if row_cells:
                rows_data.append(row_cells)
        
        tables.append({
            "table_index": index + 1,
            "rows": rows_data
        })
    return tables


def split_docx_paragraphs(paragraphs: List[str], doc_name: str) -> List[Dict[str, Any]]:
    """Chia nội dung paragraphs thành các chunks theo cấu trúc pháp luật Việt Nam.
    
    Nối các paragraphs thành một chuỗi plain text, sau đó gọi 
    VietnameseLegalSplitter để chia theo Phần > Chương > Mục > Điều > Khoản.
    
    Args:
        paragraphs: Danh sách các đoạn văn bản từ get_docx_paragraphs()
        doc_name: Tên tài liệu (dùng làm document_id trong metadata)
        
    Returns:
        Danh sách các chunks (dict) với content và metadata.
    """
    # Nối paragraphs bằng newline để tạo ra plain text hoàn chỉnh
    full_text = "\n".join(paragraphs)
    
    splitter = LLMVietnameseLegalSplitter(doc_name=doc_name)
    chunks = splitter.split_text(full_text)
    return chunks


def process_legal_documents():
    """Xử lý tất cả file .docx trong data/raw.
    
    Pipeline:
    1. Đọc paragraphs từ .docx → lưu plain text vào data/preprocessed/{name}.txt
    2. Đọc tables từ .docx → log ra console
    3. Chia chunks từ paragraphs → lưu JSON vào data/chunked/{name}.json
    """
    # Setup directories
    raw_dir = Path("data/raw")
    preprocessed_dir = Path("data/preprocessed")
    chunked_dir = Path("data/chunked")

    preprocessed_dir.mkdir(parents=True, exist_ok=True)
    chunked_dir.mkdir(parents=True, exist_ok=True)

    docx_files = list(raw_dir.glob("*.docx"))
    print(f"Found {len(docx_files)} DOCX files in {raw_dir}")

    for docx_path in docx_files:
        doc_name = docx_path.stem
        print(f"\nProcessing: {docx_path.name}")

        try:
            # 1. Trích xuất paragraphs và lưu plain text
            paragraphs = get_docx_paragraphs(docx_path)
            plain_text = "\n".join(paragraphs)

            txt_path = preprocessed_dir / f"{doc_name}.txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(plain_text)
            print(f"  - Saved {len(paragraphs)} paragraphs to: {txt_path}")

            # 2. Trích xuất tables và log thông tin
            tables = get_docx_tables(docx_path)
            print(f"  - Found {len(tables)} tables (tiêu ngữ, nơi nhận, chữ ký...)")
            for tbl in tables:
                print(f"    [Bảng {tbl['table_index']}] {len(tbl['rows'])} rows")

            # 3. Chia chunks từ paragraphs
            chunks = split_docx_paragraphs(paragraphs, doc_name)

            json_path = chunked_dir / f"{doc_name}.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(chunks, f, ensure_ascii=False, indent=2)
            print(f"  - Saved {len(chunks)} chunks to: {json_path}")

        except Exception as e:
            print(f"  - [ERROR] Failed to process {docx_path.name}: {e}")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    process_legal_documents()